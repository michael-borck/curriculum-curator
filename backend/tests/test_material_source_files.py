"""
Tests for the material source-file endpoints (Phase 3, Mode B).

- GET  /api/materials/{id}/source-files/{filename}        download
- POST /api/materials/{id}/source-files/{filename}/promote re-parse + replace

Exercises ownership (404 for someone else's material), path-traversal
rejection (400 for separators), and a promote round-trip that replaces
the material's content_json from a real DOCX source file.

Real in-memory SQLite + a tmp-dir git singleton so source-file writes
don't touch the real content_repos directory.
"""

from __future__ import annotations

import io
import uuid
from typing import TYPE_CHECKING

from docx import Document

from app.models.weekly_material import WeeklyMaterial
from app.services import git_content_service
from app.services.material_parsers.persistence import material_source_files_dir

if TYPE_CHECKING:
    from pathlib import Path

    from fastapi.testclient import TestClient
    from sqlalchemy.orm import Session

    from app.models.unit import Unit
    from app.models.user import User


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Promoted Handout", level=1)
    doc.add_paragraph("Content from the promoted source file.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _make_material_with_source(
    db: Session,
    unit: Unit,
    user_email: str,
    *,
    source_name: str,
    source_bytes: bytes,
) -> WeeklyMaterial:
    """Create a material in the DB and attach one source file in git."""
    material = WeeklyMaterial(
        id=str(uuid.uuid4()),
        unit_id=str(unit.id),
        week_number=1,
        title="Original Lecture",
        type="lecture",
        content_json={"type": "doc", "content": [{"type": "paragraph"}]},
        status="draft",
    )
    db.add(material)
    db.commit()
    db.refresh(material)

    git = git_content_service.get_git_service()
    src_dir = material_source_files_dir(material)
    git.save_binary(
        unit_id=str(unit.id),
        path=f"{src_dir}/{source_name}",
        data=source_bytes,
        user_email=user_email,
        message="seed source file",
    )
    return material


class TestDownloadSourceFile:
    def test_download_returns_bytes(
        self,
        client: TestClient,
        test_unit: Unit,
        test_user: User,
        test_db: Session,
        tmp_path: Path,
    ) -> None:
        original = git_content_service._git_service
        git_content_service._git_service = git_content_service.GitContentService(
            repos_base=str(tmp_path)
        )
        try:
            payload = _docx_bytes()
            material = _make_material_with_source(
                test_db,
                test_unit,
                test_user.email,
                source_name="lecture_01.docx",
                source_bytes=payload,
            )
            response = client.get(
                f"/api/materials/{material.id}/source-files/lecture_01.docx"
            )
            assert response.status_code == 200, response.text
            assert response.content == payload
            assert "attachment" in response.headers.get("content-disposition", "")
        finally:
            git_content_service._git_service = original

    def test_download_missing_file_404(
        self,
        client: TestClient,
        test_unit: Unit,
        test_user: User,
        test_db: Session,
        tmp_path: Path,
    ) -> None:
        original = git_content_service._git_service
        git_content_service._git_service = git_content_service.GitContentService(
            repos_base=str(tmp_path)
        )
        try:
            material = _make_material_with_source(
                test_db,
                test_unit,
                test_user.email,
                source_name="present.docx",
                source_bytes=_docx_bytes(),
            )
            response = client.get(
                f"/api/materials/{material.id}/source-files/absent.docx"
            )
            assert response.status_code == 404
        finally:
            git_content_service._git_service = original

    def test_download_unknown_material_404(self, client: TestClient) -> None:
        response = client.get(f"/api/materials/{uuid.uuid4()}/source-files/x.docx")
        assert response.status_code == 404

    def test_download_rejects_path_traversal(
        self,
        client: TestClient,
        test_unit: Unit,
        test_user: User,
        test_db: Session,
        tmp_path: Path,
    ) -> None:
        original = git_content_service._git_service
        git_content_service._git_service = git_content_service.GitContentService(
            repos_base=str(tmp_path)
        )
        try:
            material = _make_material_with_source(
                test_db,
                test_unit,
                test_user.email,
                source_name="ok.docx",
                source_bytes=_docx_bytes(),
            )
            # Encoded separators that decode to a traversal attempt. FastAPI
            # decodes %2F to "/" before our handler, so the filename arrives
            # with a separator and is rejected with 400.
            response = client.get(
                f"/api/materials/{material.id}/source-files/..%2F..%2Fsecret",
            )
            assert response.status_code == 400
        finally:
            git_content_service._git_service = original


class TestPromoteSourceFile:
    def test_promote_replaces_content_json(
        self,
        client: TestClient,
        test_unit: Unit,
        test_user: User,
        test_db: Session,
        tmp_path: Path,
    ) -> None:
        original = git_content_service._git_service
        git_content_service._git_service = git_content_service.GitContentService(
            repos_base=str(tmp_path)
        )
        try:
            material = _make_material_with_source(
                test_db,
                test_unit,
                test_user.email,
                source_name="handout.docx",
                source_bytes=_docx_bytes(),
            )
            old_content = material.content_json

            response = client.post(
                f"/api/materials/{material.id}/source-files/handout.docx/promote"
            )
            assert response.status_code == 200, response.text
            body = response.json()
            # Returned the updated material with new content_json
            assert body["id"] == str(material.id)
            assert body["contentJson"]["type"] == "doc"
            assert body["contentJson"] != old_content
            assert "promoted source file" in str(body["contentJson"]).lower()

            # Persisted in the DB
            test_db.refresh(material)
            assert material.content_json == body["contentJson"]
        finally:
            git_content_service._git_service = original

    def test_promote_missing_file_404(
        self,
        client: TestClient,
        test_unit: Unit,
        test_user: User,
        test_db: Session,
        tmp_path: Path,
    ) -> None:
        original = git_content_service._git_service
        git_content_service._git_service = git_content_service.GitContentService(
            repos_base=str(tmp_path)
        )
        try:
            material = _make_material_with_source(
                test_db,
                test_unit,
                test_user.email,
                source_name="present.docx",
                source_bytes=_docx_bytes(),
            )
            response = client.post(
                f"/api/materials/{material.id}/source-files/absent.docx/promote"
            )
            assert response.status_code == 404
        finally:
            git_content_service._git_service = original

    def test_promote_rejects_path_traversal(
        self,
        client: TestClient,
        test_unit: Unit,
        test_user: User,
        test_db: Session,
        tmp_path: Path,
    ) -> None:
        original = git_content_service._git_service
        git_content_service._git_service = git_content_service.GitContentService(
            repos_base=str(tmp_path)
        )
        try:
            material = _make_material_with_source(
                test_db,
                test_unit,
                test_user.email,
                source_name="ok.docx",
                source_bytes=_docx_bytes(),
            )
            response = client.post(
                f"/api/materials/{material.id}/source-files/..%2Fevil.docx/promote"
            )
            assert response.status_code == 400
        finally:
            git_content_service._git_service = original
