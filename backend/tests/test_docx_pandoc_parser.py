"""
Tests for the DOCX Pandoc-backed material parser.

Builds real DOCX files in memory using python-docx and runs them through
the parser. Each test that requires the Pandoc binary skips cleanly if
pandoc is not on the PATH (so CI without pandoc still passes).
"""

from __future__ import annotations

import io
from typing import Any

import pytest
from docx import Document

from app.services.material_parsers.docx_pandoc import DocxPandocParser

# Skip the entire module if pandoc isn't available — these tests exercise
# the real binary, mocking it would defeat the point
pytestmark = pytest.mark.skipif(
    not DocxPandocParser.is_available(),
    reason="pandoc binary not available",
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _save(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _node_types(content: list[dict[str, Any]]) -> list[str]:
    return [n.get("type", "") for n in content]


def _all_text(node: dict[str, Any]) -> str:
    if node.get("type") == "text":
        return str(node.get("text", ""))
    return "".join(
        _all_text(child)
        for child in node.get("content", [])
        if isinstance(child, dict)
    )


def _find_first(content: list[dict[str, Any]], node_type: str) -> dict[str, Any] | None:
    for n in content:
        if n.get("type") == node_type:
            return n
        children = n.get("content", [])
        if isinstance(children, list):
            found = _find_first(children, node_type)
            if found:
                return found
    return None


@pytest.fixture
def parser() -> DocxPandocParser:
    return DocxPandocParser()


# ---------------------------------------------------------------------------
# Builders for in-memory DOCX
# ---------------------------------------------------------------------------


def _basic_doc(title: str, body: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)
    return _save(doc)


def _doc_with_lists() -> bytes:
    doc = Document()
    doc.add_heading("List Doc", level=1)
    doc.add_paragraph("Bullets:")
    doc.add_paragraph("First bullet", style="List Bullet")
    doc.add_paragraph("Second bullet", style="List Bullet")
    doc.add_paragraph("Numbered:")
    doc.add_paragraph("First number", style="List Number")
    doc.add_paragraph("Second number", style="List Number")
    return _save(doc)


def _doc_with_table() -> bytes:
    doc = Document()
    doc.add_heading("Table Doc", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Row 1A"
    table.cell(1, 1).text = "Row 1B"
    table.cell(2, 0).text = "Row 2A"
    table.cell(2, 1).text = "Row 2B"
    return _save(doc)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestDocxBasic:
    @pytest.mark.asyncio
    async def test_heading_and_body(self, parser: DocxPandocParser) -> None:
        result = await parser.parse(
            _basic_doc("Imported Doc", "First paragraph body."),
            "doc.docx",
        )
        assert result.parser_used == "docx_pandoc"

        # Heading is the first node
        heading = _find_first(result.content_json["content"], "heading")
        assert heading is not None
        assert "Imported Doc" in _all_text(heading)

        # Body text appears as a paragraph
        paragraphs = [
            n for n in result.content_json["content"] if n.get("type") == "paragraph"
        ]
        assert any("First paragraph body" in _all_text(p) for p in paragraphs)

    @pytest.mark.asyncio
    async def test_title_extracted_from_first_heading(
        self, parser: DocxPandocParser
    ) -> None:
        result = await parser.parse(
            _basic_doc("My Title", "Body"),
            "fallback.docx",
        )
        assert result.title == "My Title"

    @pytest.mark.asyncio
    async def test_title_falls_back_to_filename(
        self, parser: DocxPandocParser
    ) -> None:
        # Doc with no headings
        doc = Document()
        doc.add_paragraph("Just a paragraph, no heading.")
        result = await parser.parse(_save(doc), "myfile.docx")
        assert result.title == "myfile"


class TestDocxLists:
    @pytest.mark.asyncio
    async def test_bullet_and_numbered_lists(
        self, parser: DocxPandocParser
    ) -> None:
        result = await parser.parse(_doc_with_lists(), "lists.docx")
        # At least one list should be present (Pandoc converts Word
        # bullet/number styles to markdown lists)
        bullet_or_ordered = [
            n
            for n in result.content_json["content"]
            if n.get("type") in ("bulletList", "orderedList")
        ]
        assert len(bullet_or_ordered) > 0
        # Bullet text should appear somewhere
        all_text = "".join(
            _all_text(n) for n in result.content_json["content"]
        )
        assert "First bullet" in all_text
        assert "First number" in all_text


class TestDocxTables:
    @pytest.mark.asyncio
    async def test_table_extracted(self, parser: DocxPandocParser) -> None:
        result = await parser.parse(_doc_with_table(), "table.docx")
        table = _find_first(result.content_json["content"], "table")
        assert table is not None
        # Cell text preserved
        assert "Header A" in _all_text(table)
        assert "Row 2B" in _all_text(table)
